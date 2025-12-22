"""DOCX document parser using python-docx."""

from typing import Optional, List
from docx import Document as DocxDocument
from docx.text.paragraph import Paragraph
from docx.table import Table

from .base import BaseParser
from .language import detect_language


class DOCXParser(BaseParser):
    """Parser for DOCX documents using python-docx."""

    def __init__(self, file_path: str):
        """Initialize DOCX parser.

        Args:
            file_path: Path to DOCX file
        """
        super().__init__(file_path)
        self._doc: Optional[DocxDocument] = None

    def _extract_text(self) -> str:
        """Extract text from DOCX.

        Returns:
            Extracted text
        """
        self._doc = DocxDocument(self.file_path)

        text_parts = []

        # Extract from paragraphs
        for paragraph in self._doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                # Preserve heading information for structure extraction
                if paragraph.style.name.startswith('Heading'):
                    level = self._get_heading_level(paragraph.style.name)
                    text_parts.append(f"[HEADING{level}] {text}")
                else:
                    text_parts.append(text)

        # Extract from tables
        for table in self._doc.tables:
            table_text = self._extract_table_text(table)
            if table_text:
                text_parts.append(f"[TABLE]\n{table_text}\n[/TABLE]")

        return "\n\n".join(text_parts)

    def _extract_table_text(self, table: Table) -> str:
        """Extract text from a table.

        Args:
            table: python-docx Table object

        Returns:
            Table text
        """
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                cells.append(cell_text)
            if any(cells):  # Skip empty rows
                rows.append(" | ".join(cells))

        return "\n".join(rows)

    def _get_heading_level(self, style_name: str) -> int:
        """Get heading level from style name.

        Args:
            style_name: Style name (e.g., 'Heading 1')

        Returns:
            Heading level (1-9)
        """
        import re
        match = re.search(r'\d+', style_name)
        if match:
            return int(match.group())
        return 1

    def _get_page_count(self) -> Optional[int]:
        """Get number of pages in DOCX.

        Note: python-docx doesn't provide direct page count.
        This is an approximation.

        Returns:
            Approximate number of pages or None
        """
        # python-docx doesn't have direct page count
        # We can approximate based on content length
        # Typical page ~500 words, ~3000 characters
        if self._doc is None:
            self._doc = DocxDocument(self.file_path)

        total_chars = sum(len(p.text) for p in self._doc.paragraphs)
        estimated_pages = max(1, total_chars // 3000)

        return estimated_pages

    def _extract_metadata(self) -> dict:
        """Extract DOCX metadata.

        Returns:
            Dictionary of metadata
        """
        if self._doc is None:
            self._doc = DocxDocument(self.file_path)

        metadata = super()._extract_metadata()

        # Add DOCX-specific metadata
        core_props = self._doc.core_properties

        metadata.update({
            "docx_title": core_props.title or "",
            "docx_author": core_props.author or "",
            "docx_subject": core_props.subject or "",
            "docx_keywords": core_props.keywords or "",
            "docx_created": str(core_props.created) if core_props.created else "",
            "docx_modified": str(core_props.modified) if core_props.modified else "",
            "docx_last_modified_by": core_props.last_modified_by or "",
            "is_docx": True,
        })

        return metadata

    def _extract_title(self, text: str) -> Optional[str]:
        """Extract title from DOCX.

        Args:
            text: DOCX text

        Returns:
            Title
        """
        # First try to get title from metadata
        if self._doc:
            metadata_title = self._doc.core_properties.title
            if metadata_title and len(metadata_title.strip()) > 3:
                return metadata_title.strip()

        # Try to find first Heading 1
        for paragraph in self._doc.paragraphs:
            if paragraph.style.name == 'Heading 1':
                title = paragraph.text.strip()
                if title:
                    return title[:200]

        # Fallback to text-based extraction
        return super()._extract_title(text)

    def _detect_language(self, text: str) -> str:
        """Detect document language.

        Args:
            text: Document text

        Returns:
            Language code
        """
        return detect_language(text)

    def get_headings(self) -> List[dict]:
        """Extract all headings from document.

        Returns:
            List of heading dictionaries
        """
        if self._doc is None:
            self._doc = DocxDocument(self.file_path)

        headings = []

        for i, paragraph in enumerate(self._doc.paragraphs):
            if paragraph.style.name.startswith('Heading'):
                level = self._get_heading_level(paragraph.style.name)
                headings.append({
                    "level": level,
                    "text": paragraph.text.strip(),
                    "index": i,
                    "style": paragraph.style.name,
                })

        return headings

    def has_tables(self) -> bool:
        """Check if document contains tables.

        Returns:
            True if document has tables
        """
        if self._doc is None:
            self._doc = DocxDocument(self.file_path)

        return len(self._doc.tables) > 0
