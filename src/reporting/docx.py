"""
DOCX Report Generator
====================

Word document report generation using python-docx.

Includes:
- Same structure as PDF
- Proper heading styles (Heading 1, 2, 3)
- Tables for findings
- Editable format for client modifications
"""

from typing import List, Dict
from io import BytesIO
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    raise ImportError("python-docx is required for Word generation. Install with: pip install python-docx")

from src.reporting.base import BaseReportGenerator


class DOCXReportGenerator(BaseReportGenerator):
    """
    Generate professional Word compliance reports.
    """

    def __init__(self):
        """Initialize DOCX generator."""
        super().__init__()

    def generate(self, findings: List[Dict], config: Dict) -> bytes:
        """
        Generate Word document report.

        Args:
            findings: List of findings
            config: Configuration dictionary

        Returns:
            bytes: DOCX file as bytes
        """
        # Create document
        doc = Document()

        # Set document properties
        doc.core_properties.title = "Compliance Gap Analysis Report"
        doc.core_properties.author = "EY Compliance Oracle"
        doc.core_properties.created = self.report_date

        # Apply custom styles
        self._apply_custom_styles(doc)

        # Cover page
        self._create_cover_page(doc, config)
        doc.add_page_break()

        # Executive Summary
        self._create_executive_summary(doc, findings, config)
        doc.add_page_break()

        # Methodology
        self._create_methodology_section(doc)
        doc.add_page_break()

        # Summary Statistics
        self._create_summary_statistics(doc, findings)
        doc.add_page_break()

        # Findings by Severity
        self._create_findings_by_severity(doc, findings)
        doc.add_page_break()

        # Findings by Domain
        self._create_findings_by_domain(doc, findings)

        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer.getvalue()

    def _apply_custom_styles(self, doc: Document):
        """Apply custom styles to document."""
        styles = doc.styles

        # Modify Heading 1
        heading1 = styles['Heading 1']
        heading1.font.name = 'Calibri'
        heading1.font.size = Pt(18)
        heading1.font.color.rgb = RGBColor(46, 46, 56)  # EY Black
        heading1.font.bold = True

        # Modify Heading 2
        heading2 = styles['Heading 2']
        heading2.font.name = 'Calibri'
        heading2.font.size = Pt(14)
        heading2.font.color.rgb = RGBColor(75, 75, 89)  # EY Dark Gray
        heading2.font.bold = True

        # Modify Heading 3
        heading3 = styles['Heading 3']
        heading3.font.name = 'Calibri'
        heading3.font.size = Pt(12)
        heading3.font.color.rgb = RGBColor(75, 75, 89)
        heading3.font.bold = True

    def _create_cover_page(self, doc: Document, config: Dict):
        """Create cover page."""
        # EY Logo (text)
        logo_para = doc.add_paragraph()
        logo_run = logo_para.add_run('EY')
        logo_run.font.size = Pt(48)
        logo_run.font.bold = True
        logo_run.font.color.rgb = RGBColor(46, 46, 56)
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add highlighting (yellow background)
        logo_para.paragraph_format.left_indent = Inches(2)
        logo_para.paragraph_format.right_indent = Inches(2)

        doc.add_paragraph()

        # Title
        title = doc.add_paragraph()
        title_run = title.add_run('Compliance Gap Analysis Report')
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Jurisdiction and domains
        jurisdiction = config.get("jurisdiction", "N/A")
        domains = config.get("selected_domains", [])
        domains_str = ", ".join(domains) if domains else "N/A"

        jurisdiction_para = doc.add_paragraph()
        jurisdiction_para.add_run('Jurisdiction: ').bold = True
        jurisdiction_para.add_run(jurisdiction)
        jurisdiction_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        domains_para = doc.add_paragraph()
        domains_para.add_run('Domains: ').bold = True
        domains_para.add_run(domains_str)
        domains_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Date
        date_para = doc.add_paragraph()
        date_para.add_run('Report Date: ').bold = True
        date_para.add_run(self.format_date())
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()
        doc.add_paragraph()

        # Confidentiality notice
        conf_para = doc.add_paragraph()
        conf_run = conf_para.add_run('CONFIDENTIAL')
        conf_run.font.bold = True
        conf_run.font.color.rgb = RGBColor(220, 53, 69)  # Red
        conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        conf_text = doc.add_paragraph(
            'This report contains confidential information intended solely for the use of the addressed client. '
            'Unauthorized disclosure, copying, or distribution is strictly prohibited.'
        )
        conf_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
        conf_text.paragraph_format.left_indent = Inches(1)
        conf_text.paragraph_format.right_indent = Inches(1)

    def _create_executive_summary(self, doc: Document, findings: List[Dict], config: Dict):
        """Create executive summary section."""
        doc.add_heading('Executive Summary', level=1)

        summary_text = self.generate_executive_summary(findings, config)

        for paragraph in summary_text.split('\n\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())

    def _create_methodology_section(self, doc: Document):
        """Create methodology section."""
        doc.add_heading('Methodology', level=1)

        methodology_text = self.generate_methodology_section()

        for paragraph in methodology_text.split('\n\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())

    def _create_summary_statistics(self, doc: Document, findings: List[Dict]):
        """Create summary statistics section."""
        doc.add_heading('Summary Statistics', level=1)

        summary = self.format_finding_summary(findings)

        # Create summary table
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Metric'
        header_cells[1].text = 'Count'

        # Data rows
        table.rows[1].cells[0].text = 'Total Findings'
        table.rows[1].cells[1].text = str(summary['total_findings'])

        table.rows[2].cells[0].text = 'Critical Priority'
        table.rows[2].cells[1].text = str(summary['by_severity']['critical'])

        table.rows[3].cells[0].text = 'High Priority'
        table.rows[3].cells[1].text = str(summary['by_severity']['high'])

        table.rows[4].cells[0].text = 'Medium Priority'
        table.rows[4].cells[1].text = str(summary['by_severity']['medium'])

        table.rows[5].cells[0].text = 'Low Priority'
        table.rows[5].cells[1].text = str(summary['by_severity']['low'])

        doc.add_paragraph()

        # Domain breakdown
        if summary['by_domain']:
            doc.add_heading('Findings by Domain', level=2)

            domain_table = doc.add_table(rows=len(summary['by_domain'])+1, cols=2)
            domain_table.style = 'Light Grid Accent 1'

            # Header
            domain_table.rows[0].cells[0].text = 'Domain'
            domain_table.rows[0].cells[1].text = 'Count'

            # Data
            for i, (domain, count) in enumerate(summary['by_domain'].items(), start=1):
                domain_table.rows[i].cells[0].text = domain
                domain_table.rows[i].cells[1].text = str(count)

    def _create_findings_by_severity(self, doc: Document, findings: List[Dict]):
        """Create findings grouped by severity."""
        doc.add_heading('Findings by Severity', level=1)

        categorized = self.categorize_by_severity(findings)

        for severity in ['critical', 'high', 'medium', 'low']:
            severity_findings = categorized.get(severity, [])

            if not severity_findings:
                continue

            severity_text = f"{severity.upper()} Priority ({len(severity_findings)} finding{'s' if len(severity_findings) != 1 else ''})"
            doc.add_heading(severity_text, level=2)

            for finding in severity_findings:
                self._create_finding_card(doc, finding)

    def _create_findings_by_domain(self, doc: Document, findings: List[Dict]):
        """Create findings grouped by domain."""
        doc.add_heading('Findings by Domain', level=1)

        categorized = self.categorize_by_domain(findings)

        for domain, domain_findings in categorized.items():
            doc.add_heading(f"{domain} ({len(domain_findings)} finding{'s' if len(domain_findings) != 1 else ''})",
                          level=2)

            for finding in domain_findings:
                self._create_finding_card(doc, finding)

    def _create_finding_card(self, doc: Document, finding: Dict):
        """Create a finding card in the document."""
        severity = finding.get('severity', 'medium')
        domain = finding.get('domain', 'Unknown')
        requirement_citation = finding.get('requirement_citation', 'N/A')
        requirement_text = finding.get('requirement_text', '')
        policy_reference = finding.get('policy_reference', 'N/A')
        gap_description = finding.get('gap_description', '')
        recommendation = finding.get('recommendation', '')
        confidence = finding.get('confidence', 0.0)

        # Header
        header = doc.add_paragraph()
        header_run = header.add_run(f"[{severity.upper()}] {domain}")
        header_run.font.bold = True
        header_run.font.size = Pt(12)

        # Requirement
        req_para = doc.add_paragraph()
        req_para.add_run('Requirement: ').bold = True
        req_para.add_run(f"{self.format_citation(requirement_citation)}\n")
        req_para.add_run(requirement_text)
        req_para.paragraph_format.left_indent = Inches(0.5)

        # Policy reference
        if policy_reference and policy_reference != 'N/A':
            policy_para = doc.add_paragraph()
            policy_para.add_run('Policy Reference: ').bold = True
            policy_para.add_run(policy_reference)

        # Gap description
        gap_para = doc.add_paragraph()
        gap_para.add_run('Gap Identified: ').bold = True
        gap_para.add_run('\n' + gap_description)

        # Recommendation
        rec_para = doc.add_paragraph()
        rec_para.add_run('💡 Recommendation: ').bold = True
        rec_para.add_run('\n' + recommendation)

        # Confidence
        conf_para = doc.add_paragraph()
        conf_para.add_run('Confidence: ').bold = True
        conf_para.add_run(self.format_confidence(confidence))

        # Spacer
        doc.add_paragraph()
